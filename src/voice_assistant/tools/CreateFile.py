import os
from agency_swarm.tools import BaseTool
from dotenv import load_dotenv
from pydantic import Field
from docx import Document
from bs4 import BeautifulSoup

from voice_assistant.config import SCRATCH_PAD_DIR
from voice_assistant.models import CreateFileResponse
from voice_assistant.utils.decorators import timeit_decorator
from voice_assistant.utils.llm_utils import get_structured_output_completion

load_dotenv()


class CreateFile(BaseTool):
    """A tool for creating a new file with generated content based on a prompt."""

    file_name: str = Field(..., description="The name of the file to be created.")
    prompt: str = Field(
        ..., description="The prompt to generate content for the new file."
    )
    format: str = Field(
        "docx", description="The format of the file content, e.g., 'docx'."
    )

    async def run(self):
        result = await create_file(self.file_name, self.prompt, self.format)
        return str(result)


@timeit_decorator
async def create_file(file_name: str, prompt: str, format: str) -> dict:
    # Ensure the file has the correct extension
    if format == "docx" and not file_name.endswith(".docx"):
        file_name += ".docx"

    file_path = os.path.join(SCRATCH_PAD_DIR, file_name)

    if os.path.exists(file_path):
        return {"status": "File already exists"}

    prompt_structure = f"""
    <purpose>
        Generate content for a new file based on the user's prompt and the file name.
    </purpose>

    <instructions>
        <instruction>Based on the user's prompt and the file name. Generate content for a new file in HTML format, which can be converted to DOCX</instruction>
        <instruction>The file name is: {file_name}</instruction>
        <instruction>Use the following prompt to generate the content: {prompt}</instruction>
    </instructions>
    """

    response = await get_structured_output_completion(
        prompt_structure, CreateFileResponse
    )

    html_content = response.file_content

    if format == "docx":
        # Parse HTML content
        soup = BeautifulSoup(html_content, "html.parser")
        # Create a Word document
        doc = Document()
        # Add content to the document
        for element in soup.find_all():
            if element.name == 'h1':
                doc.add_heading(element.get_text(), level=1)
            elif element.name == 'h2':
                doc.add_heading(element.get_text(), level=2)
            elif element.name == 'h3':
                doc.add_heading(element.get_text(), level=3)
            elif element.name == 'strong':
                p = doc.add_paragraph(element.get_text())
                p.runs[0].bold = True
            elif element.name == 'li':
                doc.add_paragraph(element.get_text(), style='List Bullet')
            elif element.name == 'p':
                doc.add_paragraph(element.get_text())
            elif element.name is None and element.strip():  # Handle plain text not wrapped in tags
                doc.add_paragraph(element.strip())
        doc.save(file_path)
    else:
        with open(file_path, "w") as f:
            f.write(html_content)

    return {"status": "File created", "file_name": response.file_name}


if __name__ == "__main__":
    import asyncio

    tool = CreateFile(file_name="test", prompt="Write a short story about a robot.", format="docx")
    print(asyncio.run(tool.run()))
