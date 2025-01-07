import os
from docx import Document
from bs4 import BeautifulSoup
import json

from agency_swarm.tools import BaseTool
from dotenv import load_dotenv
from pydantic import Field

from voice_assistant.config import SCRATCH_PAD_DIR
from voice_assistant.models import FileSelectionResponse, ModelName, CreateFileResponse
from voice_assistant.utils.decorators import timeit_decorator
from voice_assistant.utils.llm_utils import (
    get_structured_output_completion,
    parse_chat_completion,
)

load_dotenv()


class CreateProposal(BaseTool):
    """A tool for creating a proposal from a text file based on a predefined template using GPT."""

    prompt: str = Field(..., description="The prompt to identify which file to use for creating the proposal.")
    output_name: str = Field(..., description="The name of the output DOCX file.")

    async def run(self):
        result = await create_proposal(self.prompt, self.output_name)
        return str(result)


@timeit_decorator
async def create_proposal(prompt: str, output_name: str) -> dict:
    available_files = os.listdir(SCRATCH_PAD_DIR)
    available_model_map = {model.value: model.name for model in ModelName}

    file_selection_response = await get_structured_output_completion(
        create_file_selection_prompt(
            available_files, json.dumps(available_model_map), prompt
        ),
        FileSelectionResponse,
    )

    if not file_selection_response.file:
        return {"status": "No matching file found"}

    selected_file = file_selection_response.file
    file_path = os.path.join(SCRATCH_PAD_DIR, selected_file)

    with open(file_path, "r", encoding="utf-8") as f:
        file_content = f.read()

    # Create a structured prompt for GPT
    prompt_structure = f"""
    <purpose>
        Generate a proposal based on the meeting minutes provided in the text file.
    </purpose>

    <instructions>
        <instruction>Use the meeting minutes to create a proposal with the following sections: Project Overview, Scope of Work, Deliverables, Next Steps, Additional Notes.</instruction> The first page should always contain this information:
        - **Title:** Proposal for [Project Title] MVP
        - **Subtitle:** For [Client Name]
        - **Prepared & Submitted By:** Ali, Sketric Solutions
        - **Contact Number:** +13022192656
        - **Email Address:** ali@sketricsolutions.com

        <instruction>For each milestone under Scope of Work, include: Objective, Deliverables, Timeline (in weeks), and Cost (in USD).</instruction>
        <instruction>Under Deliverables, include specific items like the codebase, deployment guide, or documentation.</instruction>
        <instruction>Under Next Steps, mention client dependencies (e.g., sample data) and preparation tasks.</instruction>
        <instruction>Ensure the proposal is well-structured, coherent, and formatted professionally.</instruction>
        <instruction>Based on the user's prompt and the file name, generate content for a new file in HTML format, which can be converted to DOCX.</instruction>
    </instructions>

    <meeting-minutes>
    {file_content}
    </meeting-minutes>
    """

    # Ensure the output file has the correct extension
    if not output_name.endswith('.docx'):
        output_name += '.docx'

    output_path = os.path.join(SCRATCH_PAD_DIR, output_name)

    # Get proposal content from GPT
    response = await get_structured_output_completion(prompt_structure, CreateFileResponse)
    proposal_content = response.file_content

    # Parse HTML content
    soup = BeautifulSoup(proposal_content, "html.parser")
    # Create a Word document
    doc = Document()
    # Add content to the document
    for element in soup.find_all():
        if element.name == 'h1':
            doc.add_heading(element.get_text(), level=1)
        elif element.name == 'h2':
            doc.add_heading(element.get_text(), level=2)
        elif element.name == 'h3':
            doc.add_heading(element.get_text(), level=3)
        elif element.name == 'strong':
            p = doc.add_paragraph(element.get_text())
            p.runs[0].bold = True
        elif element.name == 'li':
            doc.add_paragraph(element.get_text(), style='List Bullet')
        elif element.name == 'p':
            doc.add_paragraph(element.get_text())
        elif element.name is None and element.strip():  # Handle plain text not wrapped in tags
            doc.add_paragraph(element.strip())
    doc.save(output_path)

    return {"status": "Proposal created", "file_name": output_name}


def create_file_selection_prompt(available_files, available_model_map, user_prompt):
    return f"""
<purpose>
    Select a file from the available files and choose the appropriate model based on the user's prompt.
</purpose>

<instructions>
    <instruction>Based on the user's prompt and the list of available files, infer which file the user wants to use for the proposal.</instruction>
    <instruction>Make sure the output the exact file name with the extension which is always .txt</instruction>
    <instruction>Also, select the most appropriate model from the available models mapping.</instruction>
    <instruction>If the user does not specify a model, default to 'base_model'.</instruction>
    <instruction>If no file matches, return an empty string for 'file'.</instruction>
</instructions>

<available-files>
    {", ".join(available_files)}
</available-files>

<available-model-map>
    {available_model_map}
</available-model-map>

<user-prompt>
    {user_prompt}
</user-prompt>
    """


if __name__ == "__main__":
    import asyncio

    tool = CreateProposal(prompt="meeting minutes for project X", output_name="proposal.docx")
    print(asyncio.run(tool.run()))
